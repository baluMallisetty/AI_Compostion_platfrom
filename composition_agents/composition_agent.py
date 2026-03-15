# composition_agent.py

import re, os
import json
from typing import Dict, List, Union, Callable

from docx import Document

# Pattern to detect {{FIELD_NAME}} style placeholders
PLACEHOLDER_PATTERN = re.compile(r"{{\s*([A-Za-z0-9_]+)\s*}}")


class CompositionAgent:
    """
    Composition Agent that:
      - Reads a Word template (.docx) with {{FIELD_ NAME}} placeholders
      - Accepts either:
          * structured JSON/dict data, or
          * free-form human text
      - Uses Phi-4 (via llm_generate callback) to fill missing fields
      - Writes data into the correct fields and saves a new .docx
    """

    def __init__(self, llm_generate: Callable[[str], str]):
        """
        llm_generate: function(prompt: str) -> str
          - This should call Phi-4 and return raw text output.
          - You will plug your existing Phi-4 client here.
        """
        self.llm_generate = llm_generate

    # -------------------------------------------------------------------------
    # 1. Discover placeholders in the template
    # -------------------------------------------------------------------------
    def _find_placeholders(self, doc: Document) -> List[str]:
        fields = set()

        # Paragraphs
        for para in doc.paragraphs:
            for match in PLACEHOLDER_PATTERN.findall(para.text):
                fields.add(match.strip())

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for match in PLACEHOLDER_PATTERN.findall(para.text):
                            fields.add(match.strip())

        return sorted(fields)

    # -------------------------------------------------------------------------
    # 2. Replace placeholders with actual values (for text + tables)
    # -------------------------------------------------------------------------
    def _replace_in_paragraph(self, paragraph, mapping: Dict[str, str]):
        if not paragraph.text:
            return

        new_text = paragraph.text

        for field, value in mapping.items():
            placeholder = "{{" + field + "}}"
            placeholder_spaced = "{{ " + field + " }}"
            new_text = new_text.replace(placeholder, value)
            new_text = new_text.replace(placeholder_spaced, value)

        # Replace runs carefully to keep formatting simple:
        for _ in range(len(paragraph.runs)):
            paragraph.runs[0].clear()
            paragraph.runs[0].text = ""

        paragraph.text = new_text

    def _replace_placeholders(self, doc: Document, mapping: Dict[str, str]):
        # Paragraphs
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, mapping)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, mapping)

    # -------------------------------------------------------------------------
    # 3. Build field mapping from JSON/dict
    # -------------------------------------------------------------------------
    def _mapping_from_json(
        self, fields: List[str], input_data: Dict
    ) -> Dict[str, str]:
        mapping: Dict[str, str] = {}
        for f in fields:
            val = input_data.get(f)
            mapping[f] = "" if val is None else str(val)
        return mapping

    # -------------------------------------------------------------------------
    # 4. Build field mapping from text using Phi-4
    #    -> single call to Phi-4 returns JSON for all fields.
    # -------------------------------------------------------------------------
    def _mapping_from_text(
        self, fields: List[str], input_text: str
    ) -> Dict[str, str]:
        fields_str = ", ".join(fields)

        prompt = f"""
You are a Composition Agent that fills a Word template.

Given:
1. A list of field names: [{fields_str}]
2. User input text (context) below.

Task:
- Infer appropriate values for EACH field.
- Return STRICTLY a valid JSON object.
- Keys MUST be EXACTLY the field names.
- Values MUST be plain strings (no newlines in keys, no extra keys).

User input text:
\"\"\"{input_text}\"\"\"
"""

        raw_output = self.llm_generate(prompt).strip()

        # Try to extract JSON from the model output
        try:
            # If the model wraps JSON in backticks or text, try to find {...}
            start = raw_output.find("{")
            end = raw_output.rfind("}")
            json_str = raw_output[start : end + 1]
            data = json.loads(json_str)
        except Exception:
            # Fallback: empty or simple mapping
            data = {}

        mapping: Dict[str, str] = {}
        for f in fields:
            val = data.get(f, "")
            mapping[f] = "" if val is None else str(val)

        return mapping

    # -------------------------------------------------------------------------
    # 5. Public API: one entrypoint for both JSON and text modes
    # -------------------------------------------------------------------------
    def compose(
        self,
        template_path: str,
        input_data: Union[str, Dict],
        output_path: str,
    ) -> Dict[str, str]:
        """
        template_path: path to .docx template
        input_data:
            - dict/JSON -> direct mapping (by field name)
            - str       -> free-form human text; Phi-4 will infer field values
        output_path: where to save the filled document

        Returns: final mapping {FIELD_NAME: value} that was written.
        """
        doc = Document(template_path)
        fields = self._find_placeholders(doc)

        if isinstance(input_data, dict):
            mapping = self._mapping_from_json(fields, input_data)
        else:
            mapping = self._mapping_from_text(fields, input_data)

        self._replace_placeholders(doc, mapping)
        doc.save(output_path)

        return mapping
    
    def compose_many(self,
                 template_path: str,
                 customers: List[Dict],
                 output_dir: str = "output"):
        """
        Generate a DOCX statement for each customer in the list.
        """
        os.makedirs(output_dir, exist_ok=True)

        results = []

        for cust in customers:
            mapping = self.mapping_from_json(cust)
            output_path = os.path.join(
                output_dir,
                f"statement_{mapping['CUSTOMER_ID']}.docx"
            )
            self.compose(
                template_path=template_path,
                input_data=cust,
                output_path=output_path
            )
            results.append({
             "customer_id": mapping["CUSTOMER_ID"],
                "file": output_path
            })
        return results

